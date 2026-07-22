from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = Path(__file__).with_name("Buku_Panduan_Pengguna_IDI_Apps.docx")
ASSETS = Path(__file__).with_name("assets")
LOGO = ROOT / "public/images/idi-logo.png"
SEEDER = ROOT / "database/seeders/RolePermissionSeeder.php"

APP_REVISION = "85fe1cb (rev 1.9.1)"
DOC_VERSION = "1.0"
DOC_DATE = "19 Juli 2026"

# compact_reference_guide + named override `idi_brand`
FONT = "Calibri"
BODY_SIZE = 11
RED = "B91C1C"
DARK_RED = "991B1B"
NAVY = "1F2937"
INK = "111827"
MUTED = "6B7280"
LIGHT_RED = "FEE2E2"
PALE_RED = "FFF7F7"
LIGHT_GRAY = "F3F4F6"
MID_GRAY = "D1D5DB"
BLUE = "1D4ED8"
GREEN = "166534"
AMBER = "92400E"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_font(run, size=None, bold=None, italic=None, color=INK, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in):
    widths = [int(round(w * 1440)) for w in widths_in]
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color="D1D5DB", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def style_table(table, header=True, font_size=8.7, header_fill=LIGHT_RED):
    set_table_borders(table)
    if header and table.rows:
        set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if header and r_idx == 0:
                set_cell_shading(cell, header_fill)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_font(run, size=font_size, bold=(header and r_idx == 0), color=INK)


def add_table(doc, headers, rows, widths, font_size=8.7, header_fill=LIGHT_RED):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)
    set_table_geometry(table, widths)
    style_table(table, font_size=font_size, header_fill=header_fill)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_field(paragraph, instruction, display=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_update_fields(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_para(doc, text="", bold=False, italic=False, color=INK, size=BODY_SIZE,
             align=None, before=0, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_label_para(doc, label, text, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    r1 = p.add_run(f"{label}: ")
    set_font(r1, size=10.3, bold=True, color=NAVY)
    r2 = p.add_run(text)
    set_font(r2, size=10.3, color=INK)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(str(item))
        set_font(r, size=10.3, color=INK)


def add_steps(doc, steps):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("w:start", "1"), ("w:numFmt", "decimal"), ("w:lvlText", "%1."), ("w:lvlJc", "left")):
        node = OxmlElement(tag)
        node.set(qn("w:val"), value)
        level.append(node)
    level_ppr = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    level_ppr.append(indent)
    level.append(level_ppr)
    abstract.append(level)
    numbering.append(abstract)

    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for step in steps:
        p = doc.add_paragraph()
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_ref = OxmlElement("w:numId")
        num_ref.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_ref)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(step)
        set_font(r, size=10.3, color=INK)


def add_callout(doc, label, text, tone="note"):
    fills = {"note": "EFF6FF", "warn": "FEF3C7", "risk": "FEE2E2", "ok": "DCFCE7"}
    colors = {"note": BLUE, "warn": AMBER, "risk": DARK_RED, "ok": GREEN}
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    # A one-cell callout is still represented as a table in DOCX; marking its
    # row as a header keeps screen-reader table audits unambiguous.
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [6.5])
    set_table_borders(table, color=fills.get(tone, LIGHT_GRAY), size="5")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fills.get(tone, LIGHT_GRAY))
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    a = p.add_run(f"{label}. ")
    set_font(a, size=10, bold=True, color=colors.get(tone, NAVY))
    b = p.add_run(text)
    set_font(b, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, filename, caption, annotations=None, width=6.35):
    path = ASSETS / filename
    if not path.exists():
        add_callout(doc, "Screenshot", "Perlu verifikasi melalui pengujian aplikasi; berkas gambar tidak tersedia.", "warn")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", caption)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(4)
    cr = cap.add_run(caption)
    set_font(cr, size=8.5, italic=True, color=MUTED)
    if annotations:
        add_callout(doc, "Anotasi", " ".join(f"[{i + 1}] {item}" for i, item in enumerate(annotations)), "note")


def add_heading(doc, number, title, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(f"{number} {title}")
    return p


def feature_block(doc, *, title, purpose, access, denied, basis, prerequisites,
                  steps, expected, errors, risks, screenshot, notes):
    add_para(doc, title, bold=True, color=DARK_RED, size=11.3, before=8, after=4, keep=True)
    add_label_para(doc, "Tujuan", purpose)
    add_label_para(doc, "Role yang dapat mengakses", access)
    add_label_para(doc, "Role yang tidak dapat mengakses", denied)
    add_label_para(doc, "Dasar permission/aturan", basis)
    add_label_para(doc, "Prasyarat", prerequisites)
    add_para(doc, "Langkah penggunaan", bold=True, color=NAVY, size=10.3, after=2, keep=True)
    add_steps(doc, steps)
    add_label_para(doc, "Hasil yang diharapkan", expected)
    add_label_para(doc, "Validasi/error yang mungkin muncul", errors)
    add_label_para(doc, "Dampak atau risiko", risks)
    if screenshot:
        add_label_para(doc, "Screenshot beranotasi", screenshot)
    else:
        add_label_para(doc, "Screenshot beranotasi", "Perlu verifikasi melalui pengujian aplikasi.")
    add_label_para(doc, "Catatan khusus", notes, after=8)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, DARK_RED, 18, 10),
        2: (13, RED, 14, 7),
        3: (12, NAVY, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        st = styles[style_name]
        st.font.name = FONT
        st.font.size = Pt(10.3)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def configure_section(section, landscape=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_running_header_footer(section, short_title="Buku Panduan Pengguna IDI Apps"):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(short_title)
    set_font(r, size=8.5, bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    left = fp.add_run(f"Versi {DOC_VERSION}  •  ")
    set_font(left, size=8, color=MUTED)
    add_field(fp, "PAGE", "1")
    for run in fp.runs:
        set_font(run, size=8, color=MUTED)


def parse_permissions():
    text = SEEDER.read_text(encoding="utf-8")
    master_match = re.search(r"\$permissions\s*=\s*\[(.*?)\];", text, re.S)
    permissions = re.findall(r"'([^']+)'", master_match.group(1))
    role_map = {"admin": list(permissions), "superadmin": list(permissions)}
    for var, role in (("sekretaris", "sekretaris"), ("bendahara", "bendahara"), ("ketua", "ketua"), ("anggota", "anggota")):
        match = re.search(rf"\${var}->syncPermissions\(\[(.*?)\]\);", text, re.S)
        role_map[role] = re.findall(r"'([^']+)'", match.group(1)) if match else []
    give = re.search(r"\$bendahara->givePermissionTo\(\[(.*?)\]\);", text, re.S)
    if give:
        role_map["bendahara"].extend(re.findall(r"'([^']+)'", give.group(1)))
    return permissions, role_map


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build():
    permissions, role_map = parse_permissions()
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], landscape=False)
    doc.sections[0].different_first_page_header_footer = True
    add_running_header_footer(doc.sections[0])

    core = doc.core_properties
    core.title = "Buku Panduan Pengguna IDI Apps"
    core.subject = "Panduan operasional, matriks akses, FAQ, dan troubleshooting"
    core.author = "IDI Cabang Purwakarta"
    core.keywords = "IDI Apps, panduan pengguna, role, permission, QA"
    core.comments = "Disusun berdasarkan audit source code dan pengujian aplikasi dengan data fiktif."
    core.created = datetime(2026, 7, 19)
    core.modified = datetime(2026, 7, 19)
    set_update_fields(doc)

    # Cover — editorial_cover named override.
    add_para(doc, "PANDUAN OPERASIONAL RESMI", bold=True, color=RED, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=20, after=18)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic = p.add_run().add_picture(str(LOGO), width=Inches(1.45))
        pic._inline.docPr.set("descr", "Logo Ikatan Dokter Indonesia")
        p.paragraph_format.space_after = Pt(20)
    add_para(doc, "BUKU PANDUAN\nPENGGUNA IDI APPS", bold=True, color=NAVY, size=28,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "Panduan fitur, alur lintas role, kontrol akses, dan penanganan masalah",
             color=MUTED, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=42)
    cover_rows = [
        ("Revisi aplikasi", APP_REVISION),
        ("Versi dokumen", DOC_VERSION),
        ("Tanggal dokumen", DOC_DATE),
        ("Status", "Final terkendali — dengan item verifikasi manual"),
        ("Pemilik dokumen", "IDI Cabang Purwakarta (berdasarkan fallback identitas aplikasi)"),
    ]
    add_table(doc, ["Informasi", "Nilai"], cover_rows, [1.65, 4.85], font_size=9.5)
    add_para(doc, "Dokumen ini tidak memuat kredensial, data anggota asli, atau data finansial nyata.",
             italic=True, color=MUTED, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, before=24)
    add_page_break(doc)

    add_heading(doc, "1", "Kontrol Dokumen", 1)
    add_heading(doc, "1.1", "Riwayat Revisi", 2)
    add_table(doc, ["Versi", "Tanggal", "Status", "Perubahan", "Penyusun"], [
        ("1.0", DOC_DATE, "Final terkendali", "Dokumentasi baseline source saat ini; 217 test otomatis lulus; screenshot memakai data fiktif.", "System Analyst / Technical Writer / QA Documentation"),
    ], [0.55, 0.9, 1.05, 3.0, 1.0], font_size=8.5)
    add_heading(doc, "1.2", "Dasar dan Batas Analisis", 2)
    add_para(doc, "Sumber kebenaran yang diaudit mencakup route, seeder role/permission, policy, controller, service, request validation, model, migration, menu frontend, halaman React, dan test Feature. Audit mencatat 223 route terdaftar, 68 controller, 51 service, 92 permission default, dan 43 berkas test Feature.")
    add_bullets(doc, [
        "Pengujian otomatis: 217 test lulus dengan 1.263 assertion.",
        "Pengujian UI dilakukan pada instance lokal terisolasi menggunakan akun dan data fiktif.",
        "Tindakan destruktif (restore dan reset) tidak dijalankan melalui UI; klaimnya dibatasi pada source dan test otomatis.",
        "Source mengalami perubahan kerja lokal selama audit; baseline aplikasi dinyatakan dengan revisi source dan temuan dari berkas yang tersedia pada tanggal dokumen.",
    ])
    add_callout(doc, "Konvensi bukti", "Label “Perlu verifikasi melalui pengujian aplikasi” berarti fitur dapat ditelusuri pada source tetapi tidak diuji tuntas melalui UI dengan kondisi data target.", "warn")

    add_heading(doc, "1.3", "Daftar Isi Otomatis", 2)
    toc = doc.add_paragraph()
    toc.paragraph_format.space_after = Pt(12)
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "Klik kanan dan pilih Update Field jika daftar belum diperbarui.")
    add_para(doc, "Daftar isi menggunakan field otomatis Word. Nomor halaman akan diperbarui saat dokumen dibuka dan field diperbarui.", italic=True, color=MUTED, size=8.5)
    add_para(doc, "Ringkasan struktur", bold=True, color=NAVY, size=10.3, before=6, after=4, keep=True)
    add_table(doc, ["Bab 1–8", "Bab 9–16"], [
        ("1 Kontrol Dokumen; 2 Pendahuluan", "9 Tindakan Berisiko; 10 Catatan Implementasi"),
        ("3 Gambaran & Navigasi; 4 Role & Keamanan Akses", "11 Troubleshooting & FAQ; 12 Verifikasi Manual"),
        ("5 Matriks Role; 6 Panduan Per Modul", "13 Lampiran Permission; 14 Kamus Istilah"),
        ("7 Alur Lintas Role; 8 Status Workflow", "15 Catatan Role Bidang; 16 Checklist QA"),
    ], [3.25, 3.25], font_size=8.5)

    chapter_two = add_heading(doc, "2", "Pendahuluan", 1)
    chapter_two.paragraph_format.page_break_before = True
    add_heading(doc, "2.1", "Tujuan", 2)
    add_para(doc, "Buku ini membantu pengguna menjalankan IDI Apps secara konsisten sesuai role, status data, scope divisi, kepemilikan, PIC/assignee, dan kebutuhan approval yang benar-benar diterapkan pada source saat ini.")
    add_heading(doc, "2.2", "Ruang Lingkup", 2)
    add_bullets(doc, [
        "Autentikasi, profil, dashboard, sekretariat, pengurus, anggota, program kerja, iuran, kas/transaksi, laporan, audit/approval, pengaturan, dan verifikasi publik.",
        "Operasi harian, validasi umum, hasil yang diharapkan, risiko, troubleshooting, FAQ, permission, dan workflow status.",
        "Tidak membahas konfigurasi server, deployment, environment variable, database server, atau infrastruktur.",
    ])
    add_heading(doc, "2.3", "Istilah dan Konvensi", 2)
    add_table(doc, ["Simbol/Istilah", "Arti"], [
        ("✓", "Dapat melakukan aksi berdasarkan permission default dan aturan status/scope."),
        ("L", "Dapat melihat; belum tentu dapat mengubah."),
        ("A", "Dapat melakukan setelah atau sebagai bagian dari approval."),
        ("—", "Tidak dapat melakukan berdasarkan default seeder."),
        ("Bidang", "Istilah aman untuk profil kerja berbasis role anggota dan scope divisi; bukan role literal."),
        ("Catatan implementasi", "Perbedaan antara UI, route, policy, service, atau test yang perlu diperhatikan."),
    ], [1.25, 5.25], font_size=9)

    add_heading(doc, "3", "Gambaran Aplikasi dan Navigasi", 1)
    add_para(doc, "Navigasi utama menggunakan menu berbasis permission gabungan (permission dari role ditambah permission individual). Menu yang terlihat bukan pengganti kontrol backend: URL langsung tetap diperiksa oleh middleware/policy.")
    add_figure(doc, "02-dashboard-admin.png", "Gambar 1. Dashboard Admin dengan navigasi dan ringkasan operasional.", [
        "Sidebar kiri menampilkan modul yang lolos filter permission.",
        "Area atas menunjukkan organisasi dan pengguna aktif.",
        "Kartu dan quick action merangkum pekerjaan harian.",
    ])
    add_callout(doc, "Catatan implementasi", "Route Dashboard hanya memerlukan autentikasi dan email terverifikasi. Pengujian role anggota memperlihatkan ringkasan finansial dan tombol quick action, meskipun URL transaksi kemudian ditolak 403. Perlakukan Dashboard sebagai informasi internal dan evaluasi kebutuhan pembatasan lebih lanjut.", "risk")
    add_heading(doc, "3.1", "Alur Navigasi Umum", 2)
    add_steps(doc, [
        "Masuk melalui Portal Internal.",
        "Pilih modul pada sidebar; buka sub-menu dengan tombol plus.",
        "Gunakan filter, pencarian, dan pagination pada halaman daftar.",
        "Buka detail sebelum menjalankan aksi yang bergantung pada status.",
        "Periksa pesan berhasil/gagal; jangan mengulang submit bila proses masih berjalan.",
        "Keluar melalui menu pengguna setelah pekerjaan selesai.",
    ])

    add_heading(doc, "4", "Role dan Prinsip Keamanan Akses", 1)
    add_heading(doc, "4.1", "Role yang Benar-benar Tersedia", 2)
    add_table(doc, ["Role literal", "Keterangan"], [
        ("superadmin", "Seluruh 92 permission default; memiliki proteksi khusus saat memberi role superadmin melalui assignment organisasi."),
        ("admin", "Seluruh 92 permission default; satu-satunya role selain superadmin yang lolos group route Pengaturan."),
        ("ketua", "Review/approval program kerja, approval void, akses sekretariat/pengurus, dan tampilan finansial tertentu."),
        ("sekretaris", "Operasional sekretariat, pengurus, dan program kerja dalam scope."),
        ("bendahara", "Operasional iuran/transaksi, audit lihat, periode lihat, serta budget program dalam scope."),
        ("anggota", "Akses program kerja dalam scope divisi/PIC/assignee, progres, dan dokumen."),
    ], [1.25, 5.25], font_size=9)
    add_callout(doc, "Penamaan aman", "Gunakan “Bidang (profil kerja: anggota + scope divisi/PIC/penugasan)” pada panduan. Jangan menulis “role bidang”, karena seeder tidak membuat role literal tersebut. Struktur organisasi memiliki unit/divisi dan jabatan seperti Ketua Bidang, tetapi role portal ditetapkan terpisah.", "note")

    add_heading(doc, "4.2", "Permission Role dan Permission Individual", 2)
    add_bullets(doc, [
        "Permission role diwariskan dari role dan menjadi baseline default.",
        "Permission individual bersifat tambahan. Sinkronisasi permission pengguna hanya mengubah permission langsung; tidak mencabut permission yang masih diwariskan dari role.",
        "Untuk mencabut kemampuan yang diwariskan, ubah permission role atau pindahkan pengguna ke role lain; uji kembali URL langsung.",
        "Perubahan role/permission dicatat ke activity log oleh service pengaturan akses.",
        "Group route Pengaturan tetap mensyaratkan role admin atau superadmin. Memberi settings.view secara individual kepada role lain tidak membuka halaman Pengaturan.",
    ])
    add_callout(doc, "Risiko kontrol akses", "Controller pengelolaan role membandingkan nama “Admin” dengan huruf besar, sedangkan seeder memakai “admin” huruf kecil. Perlindungan hapus role Admin dan pencegahan penurunan role diri sendiri perlu diverifikasi/dikoreksi sebelum diandalkan.", "risk")

    add_heading(doc, "4.3", "Menu yang Terlihat per Role", 2)
    role_rows = [
        ("Admin", "Dashboard; Sekretariat; Anggota; Program Kerja; Iuran; Kas/Transaksi; Laporan; Audit; Pengaturan"),
        ("Ketua", "Dashboard; Sekretariat; Program Kerja; Iuran; Kas/Transaksi; Laporan (Tutup Buku); Audit; Pengaturan terlihat tetapi ditolak 403"),
        ("Sekretaris", "Dashboard; Sekretariat; Program Kerja"),
        ("Bendahara", "Dashboard; Sekretariat (Tanda Tangan); Program Kerja; Iuran; Kas/Transaksi; Laporan (Tutup Buku); Audit; Pengaturan terlihat tetapi ditolak 403"),
        ("Bidang/anggota", "Dashboard; Sekretariat (Tanda Tangan); Program Kerja"),
    ]
    add_table(doc, ["Profil", "Menu hasil pengujian UI"], role_rows, [1.25, 5.25], font_size=8.8)
    for file_name, caption, note in [
        ("22-role-ketua-menu.png", "Gambar 2. Menu role Ketua.", "Pengaturan tampak tetapi route menolak 403."),
        ("23-role-sekretaris-menu.png", "Gambar 3. Menu role Sekretaris.", "Terfokus pada Sekretariat dan Program Kerja."),
        ("24-role-bendahara-menu.png", "Gambar 4. Menu role Bendahara.", "Tutup Buku terlihat; rekap iuran tidak tersedia secara default."),
        ("25-role-bidang-menu.png", "Gambar 5. Menu profil kerja Bidang/anggota.", "Sekretariat muncul karena Tanda Tangan tidak memiliki permission menu khusus."),
    ]:
        add_figure(doc, file_name, caption, [note])

    # Matrix in landscape named override.
    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(landscape, landscape=True)
    add_running_header_footer(landscape, "IDI Apps • Matriks Akses")
    add_heading(doc, "5", "Matriks Role versus Modul", 1)
    matrix = [
        ("Dashboard", "L", "L", "L", "L", "L", "auth + verified", "Semua role melihat ringkasan; catatan eksposur finansial."),
        ("Board Sekretariat", "✓", "L", "✓", "—", "—", "secretariat.view", "Bidang hanya melihat parent karena Tanda Tangan."),
        ("Pengurus/Organisasi", "✓", "✓", "✓", "L", "—", "organization.*", "Periode ended menjadi read-only."),
        ("Surat", "✓", "A", "✓", "—", "—", "letters.*", "Ketua dapat finalisasi/PDF, tidak membuat default."),
        ("Tanda Tangan", "A", "A", "A", "A", "A", "auth + verified; signer member", "Inbox tanpa permission khusus; tanda tangan hanya milik signer."),
        ("Agenda", "✓", "L", "✓", "—", "—", "agenda.view/manage", "Pengujian UI perlu diulang pada lingkungan target."),
        ("Template & Penomoran", "✓", "—", "✓", "—", "—", "templates.manage; numbering.manage", "Nomor baru dikunci saat finalisasi."),
        ("Arsip Sekretariat", "✓", "✓", "✓", "—", "—", "secretariat.view", "Route unggah memakai view, bukan manage terpisah."),
        ("Data Anggota", "✓", "—", "—", "—", "—", "members.*", "Default seeder hanya Admin/Superadmin."),
        ("Import/Export Anggota", "✓", "—", "—", "—", "—", "members.import/export/resolve_import", "Konflik perlu resolve_import."),
        ("Program Kerja - lihat", "✓", "L", "L", "L", "L", "view atau scope", "Ketua global; lainnya scope divisi/PIC/assignee."),
        ("Program - buat/edit/submit", "✓", "✓", "✓", "—", "—", "create/update/submit", "Hanya draft/revision; harus dalam scope."),
        ("Program - review/approve", "✓", "A", "—", "—", "—", "review/approve/reject", "Tidak boleh self-review/self-approve."),
        ("Tugas/Gantt/dependensi", "✓", "✓", "✓", "L", "L", "manage_tasks atau assignment", "PIC/assignee dapat progres pada status tertentu."),
        ("Budget Program", "✓", "✓", "—", "✓", "—", "manage_budget", "Terkunci setelah completed/evaluated/terminal."),
        ("Iuran Pembayaran", "✓", "L/A", "—", "✓", "—", "dues.*", "Ketua approve void; Bendahara request void."),
        ("Rekap & Ekspor Iuran", "✓", "—", "—", "—", "—", "dues.recap.view/export", "Bendahara tidak mendapat default permission."),
        ("Kas/Transaksi", "✓", "L/A", "—", "✓", "—", "transactions.*", "Nilai finansial posted immutable; void melalui approval."),
        ("Laporan Kas/Resume", "✓", "—", "—", "—", "—", "reports.*", "Default role lain tidak punya reports.*."),
        ("Tutup Buku", "✓", "L", "—", "L", "—", "finance.period.*", "Reopen punya permission/policy tetapi tidak ada route UI."),
        ("Audit & Approval", "✓", "A", "—", "L", "—", "activity.view / approve", "Pemohon tidak boleh approve permintaan sendiri."),
        ("Pengaturan", "✓", "—", "—", "—", "—", "role admin|superadmin", "Menu Ketua/Bendahara terlihat namun route 403."),
        ("Verifikasi Publik", "L", "L", "L", "L", "L", "route publik", "Tidak memerlukan login; tampilkan metadata minimum."),
    ]
    add_table(doc,
              ["Modul/Fitur", "Admin", "Ketua", "Sekr.", "Bend.", "Bidang", "Dasar permission/aturan", "Catatan batasan"],
              matrix, [1.45, 0.45, 0.5, 0.5, 0.5, 0.55, 2.25, 2.8], font_size=7.4)
    add_callout(doc, "Cara membaca", "Matriks menunjukkan default seeder. Permission individual dapat menambah kemampuan, tetapi policy status/scope dan role guard tetap berlaku.", "note")

    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(portrait, landscape=False)
    add_running_header_footer(portrait)

    add_heading(doc, "6", "Panduan Per Modul", 1)
    add_heading(doc, "6.1", "Autentikasi dan Profil", 2)
    add_figure(doc, "01-login.png", "Gambar 6. Portal login internal.", [
        "Masukkan email akun aktif.", "Gunakan Lupa kata sandi untuk meminta tautan reset.", "Registrasi publik dinonaktifkan secara default."
    ])
    feature_block(doc,
        title="Login",
        purpose="Mengautentikasi akun aktif dan membuat sesi internal.",
        access="Semua pengguna dengan akun aktif.",
        denied="Pengunjung tanpa akun, akun nonaktif, atau kredensial salah.",
        basis="Route guest; LoginRequest menambah syarat is_active=true.",
        prerequisites="Email, kata sandi, dan akun aktif.",
        steps=["Buka Portal Internal.", "Isi Email dan Kata Sandi.", "Opsional: pilih Ingat saya pada perangkat pribadi.", "Pilih Masuk."],
        expected="Dialihkan ke Dashboard.",
        errors="Kredensial tidak cocok; terlalu banyak percobaan memicu rate limit; akun nonaktif ditolak.",
        risks="Jangan gunakan Ingat saya pada perangkat bersama; jangan membagikan kata sandi.",
        screenshot="Gambar 6.",
        notes="Path login dikonfigurasi sebagai portal/akses-internal; registrasi publik tidak tersedia secara default."
    )
    feature_block(doc,
        title="Lupa/Reset Password dan Verifikasi Email",
        purpose="Memulihkan akses dan memastikan email terverifikasi sebelum membuka modul internal.",
        access="Pengguna yang menguasai email akun.",
        denied="Token reset/verifikasi invalid atau kedaluwarsa.",
        basis="Route password.request/reset dan verification.notice/verify; modul aplikasi memakai middleware verified.",
        prerequisites="Alamat email valid dan kanal email aplikasi berfungsi.",
        steps=["Pilih Lupa kata sandi.", "Kirim email akun.", "Buka tautan reset dari email.", "Masukkan kata sandi baru dan konfirmasi.", "Jika diminta, selesaikan verifikasi email."],
        expected="Kata sandi diperbarui dan akun dapat masuk; akses modul terbuka setelah email verified.",
        errors="Email tidak ditemukan, token invalid, konfirmasi tidak sama, atau pengiriman email gagal.",
        risks="Jangan meneruskan token reset; Admin reset memakai minimal 8 karakter.",
        screenshot=None,
        notes="Alur route dan test otomatis lulus. Pengiriman email nyata perlu verifikasi melalui pengujian aplikasi."
    )
    feature_block(doc,
        title="Profil Pengguna",
        purpose="Memperbarui nama/email, kata sandi, atau menghapus akun sendiri.",
        access="Semua pengguna terautentikasi dan terverifikasi.",
        denied="Pengunjung atau pengguna yang tidak memberikan kata sandi benar saat penghapusan.",
        basis="profile.edit/update/destroy dan password.update.",
        prerequisites="Sesi aktif; kata sandi saat ini untuk perubahan kata sandi/penghapusan.",
        steps=["Buka menu pengguna > Profil.", "Ubah informasi yang diperlukan.", "Simpan dan periksa pesan hasil.", "Untuk hapus akun, masukkan kata sandi dan konfirmasi."],
        expected="Profil diperbarui atau akun dihapus dan sesi berakhir.",
        errors="Email duplikat, kata sandi saat ini salah, atau validasi form gagal.",
        risks="Penghapusan akun berpotensi memutus kepemilikan aktivitas; konsultasikan Admin.",
        screenshot=None,
        notes="Test Profile mencakup update dan penghapusan dengan validasi kata sandi."
    )

    add_heading(doc, "6.2", "Dashboard", 2)
    feature_block(doc,
        title="Ringkasan Dashboard",
        purpose="Melihat saldo, pemasukan/pengeluaran, iuran, agenda, transaksi, dan surat terbaru.",
        access="Semua role yang login dan verified.",
        denied="Pengunjung dan akun belum verified.",
        basis="Route dashboard hanya auth+verified; tidak ada permission khusus.",
        prerequisites="Periode bulan dipilih; data transaksi/iuran tersedia.",
        steps=["Buka Dashboard.", "Pilih bulan.", "Baca kartu metrik dan grafik.", "Gunakan quick action bila role mengizinkan."],
        expected="Metrik bulan dan ringkasan organisasi tampil.",
        errors="Data kosong menghasilkan nilai nol/empty state; quick action dapat berakhir 403 jika role tidak berizin.",
        risks="Ringkasan finansial terlihat pada role anggota; evaluasi kebijakan kerahasiaan internal.",
        screenshot="Gambar 1 dan Gambar 5.",
        notes="Quick action bukan bukti izin backend."
    )

    add_heading(doc, "6.3", "Sekretariat — Board dan Pengurus", 2)
    add_figure(doc, "03-secretariat-board.png", "Gambar 7. Board Sekretariat.", ["Statistik draft/final/arsip.", "Surat terbaru.", "Pintasan buat surat dan unggah arsip."])
    add_figure(doc, "16-organization.png", "Gambar 8. Halaman Pengurus/struktur organisasi.", ["Pemilih periode.", "Tab struktur dan assignment.", "Status periode mengendalikan edit/read-only."])
    feature_block(doc,
        title="Board Sekretariat",
        purpose="Memantau surat, arsip, template, dan agenda dari satu halaman.",
        access="Admin, Ketua (lihat), Sekretaris.",
        denied="Bendahara dan Bidang untuk route Board.",
        basis="secretariat.view.",
        prerequisites="Sesi verified dan permission.",
        steps=["Buka Sekretariat > Board.", "Periksa statistik.", "Buka surat terbaru atau pintasan kerja."],
        expected="Ringkasan sekretariat tampil.",
        errors="403 jika secretariat.view tidak tersedia.",
        risks="Pintasan dapat memerlukan permission tambahan.",
        screenshot="Gambar 7.",
        notes="Parent Sekretariat dapat terlihat untuk Bidang karena child Tanda Tangan tidak berpermission."
    )
    feature_block(doc,
        title="Periode, Struktur, Posisi, dan Assignment Pengurus",
        purpose="Menyiapkan struktur organisasi per periode dan menyinkronkan role/divisi akun saat assignment aktif.",
        access="Admin, Ketua, Sekretaris untuk kelola; Bendahara untuk lihat/riwayat.",
        denied="Bidang secara default.",
        basis="organization.period.*, organization.structure.manage, organization.assignment.manage/replace; policy read-only.",
        prerequisites="Master divisi/posisi, member aktif, tanggal valid, slot belum terisi, dan role portal dipilih.",
        steps=["Buat periode draft.", "Susun unit dan posisi wajib.", "Tambah assignment member dan role portal.", "Jalankan pemeriksaan kesiapan.", "Publish lalu aktifkan periode.", "Gunakan Ganti/Akhiri untuk perubahan pengurus."],
        expected="Periode aktif tunggal, struktur terbentuk, assignment aktif, role/divisi akun tersinkron, audit tercatat.",
        errors="Siklus struktur, slot terisi, member ganda, periode read-only, tanggal invalid, atau role superadmin oleh non-superadmin.",
        risks="Aktivasi/replace/end mengubah akses akun; review ringkasan sebelum konfirmasi.",
        screenshot="Gambar 8.",
        notes="Periode ended/archived bersifat historis/read-only. Assignment draft tidak mengubah akses sampai aktivasi."
    )

    add_heading(doc, "6.4", "Sekretariat — Surat", 2)
    add_figure(doc, "04-letters-list.png", "Gambar 9. Daftar Surat.", ["Filter pencarian/status/template.", "Aksi Buat Surat.", "Tabel status draft/final."])
    add_figure(doc, "18-letter-detail.png", "Gambar 10. Detail Surat Final.", ["Metadata dan status.", "Preview/unduh PDF.", "Lampiran dan versi."])
    feature_block(doc,
        title="Draft, Edit, Builder, Layout, dan Lampiran Surat",
        purpose="Menyusun surat dari template hingga siap difinalisasi.",
        access="Admin dan Sekretaris; pengguna dengan letters.create/update individual.",
        denied="Ketua untuk pembuatan/edit default; Bendahara dan Bidang.",
        basis="letters.create, letters.update; max 3 signer; max 5 lampiran, masing-masing 10 MB, tipe terbatas.",
        prerequisites="Template aktif, metadata surat, isi, penerima, dan penandatangan sesuai kebutuhan.",
        steps=["Buka Surat > Buat Surat.", "Pilih template dan isi metadata.", "Atur isi, penandatangan, QR, serta layout/builder.", "Simpan Draft.", "Tambahkan lampiran pada detail surat.", "Tinjau preview."],
        expected="Surat berstatus draft dengan layout/blocks tersimpan.",
        errors="Template/member signer tidak ditemukan; file melewati ukuran/tipe; field wajib finalisasi belum lengkap.",
        risks="Perubahan layout dapat menggeser tampilan PDF; cek preview setiap revisi.",
        screenshot="Gambar 9 dan Gambar 10.",
        notes="Draft dapat mengambil struktur template ketika layout tidak dikirim; diverifikasi test otomatis."
    )
    feature_block(doc,
        title="Penomoran, Finalisasi, PDF, Versi, Arsip, dan Pencabutan",
        purpose="Mengunci nomor, menghasilkan dokumen final, menyimpan versi, dan mengelola siklus arsip/cabut.",
        access="Admin; Sekretaris; Ketua untuk finalisasi/PDF/versi. Pencabutan default hanya Admin.",
        denied="Bendahara dan Bidang.",
        basis="numbering.manage; letters.finalize/export_pdf/versions.view/revoke; status draft/finalized/archived.",
        prerequisites="Metadata final lengkap, template/format nomor valid, signer sesuai, dan nomor belum dipakai.",
        steps=["Gunakan Generate untuk pratinjau nomor.", "Tinjau surat dan signer.", "Pilih Finalisasi.", "Buka PDF/Preview dan versi.", "Arsipkan surat final bila diperlukan.", "Untuk membatalkan keabsahan, gunakan Cabut Surat dan isi alasan."],
        expected="Nomor unik dikunci saat finalisasi, versi bertambah, PDF tersedia, hash publik dibuat; surat cabut ditandai tidak valid.",
        errors="Hanya surat final dapat diarsipkan/regenerate PDF; nomor duplikat; layout kosong; alasan pencabutan kurang.",
        risks="Finalisasi membuat snapshot dan nomor resmi; pencabutan memengaruhi verifikasi publik.",
        screenshot="Gambar 10.",
        notes="Generate nomor hanya pratinjau dan tidak menaikkan sequence; test mengonfirmasi nomor pertama yang belum dipakai dipilih saat finalisasi."
    )

    add_heading(doc, "6.5", "Tanda Tangan dan Verifikasi Publik", 2)
    add_figure(doc, "06-signatures.png", "Gambar 11. Inbox Tanda Tangan.", ["Daftar permintaan milik signer.", "Status surat/signature.", "Aksi tanda tangan sesuai kepemilikan."])
    add_figure(doc, "21-public-verification.png", "Gambar 12. Verifikasi publik dokumen surat.", ["Status valid/cabut.", "Metadata minimum surat.", "Versi dokumen dan tautan PDF bila tersedia."])
    feature_block(doc,
        title="Persiapan dan Tanda Tangan Surat",
        purpose="Mencatat signer dan melengkapi metadata QR/tanda tangan surat final.",
        access="Route inbox untuk semua pengguna verified; aksi sign hanya member yang ditetapkan sebagai signer.",
        denied="Pengguna yang bukan signer, surat belum final, surat dicabut, atau permintaan dicabut.",
        basis="Route tanpa permission khusus; controller memeriksa signer_member_id, status finalized/archived, is_revoked, revoked_at.",
        prerequisites="Akun terhubung ke member signer dan surat sudah final.",
        steps=["Buka Sekretariat > Tanda Tangan.", "Pilih permintaan milik Anda.", "Periksa surat dan status.", "Pilih Tanda Tangan."],
        expected="Signature tercatat; metadata barcode/QR diperbarui ketika semua signer selesai.",
        errors="403 bukan signer; 422 surat belum final/sudah dicabut/permintaan dicabut.",
        risks="Inbox dapat terlihat bagi semua role verified; jangan mengandalkan menu sebagai pembatas.",
        screenshot="Gambar 11.",
        notes="Test mencakup multi-signer, signer stale, dan penyelesaian metadata barcode."
    )
    feature_block(doc,
        title="Verifikasi Publik Surat/Dokumen",
        purpose="Memungkinkan pihak luar memeriksa keabsahan surat melalui hash atau signature.",
        access="Publik tanpa login.",
        denied="Hash/signature tidak ditemukan atau dokumen tidak tersedia untuk download.",
        basis="letters.verify, letters.verify.download, letters.signature.verify.",
        prerequisites="Surat final memiliki public_hash/signature.",
        steps=["Buka tautan/QR verifikasi.", "Periksa status VALID/DICABUT dan metadata.", "Cocokkan nomor, tanggal, perihal, signer, dan versi.", "Unduh PDF hanya bila tautan tersedia."],
        expected="Status dan metadata minimum tampil tanpa login.",
        errors="404 untuk hash tidak dikenal; download dapat gagal bila berkas tidak ada.",
        risks="Jangan menaruh data sensitif pada perihal atau metadata publik.",
        screenshot="Gambar 12.",
        notes="Pengujian UI memakai surat fiktif berstatus valid. Skenario cabut perlu verifikasi manual."
    )

    add_heading(doc, "6.6", "Agenda, Template Surat, dan Arsip", 2)
    add_figure(doc, "19-letter-templates.png", "Gambar 13. Template Surat.", ["Daftar template aktif.", "Aksi builder/layout.", "Pengaturan signer dan nomor."])
    add_figure(doc, "20-archive.png", "Gambar 14. Arsip Dokumen Sekretariat.", ["Filter arsip.", "Unggah dokumen.", "Preview dan download."])
    feature_block(doc,
        title="Agenda",
        purpose="Mencatat agenda internal/eksternal, status, waktu, lokasi, PIC, dan lampiran.",
        access="Admin dan Sekretaris kelola; Ketua lihat.",
        denied="Bendahara dan Bidang.",
        basis="agenda.view dan agenda.manage; tipe internal/external; status planned/done/cancelled.",
        prerequisites="Judul, status, waktu mulai; waktu selesai tidak boleh sebelum mulai.",
        steps=["Buka Agenda.", "Pilih Tambah Agenda.", "Isi tipe, status, waktu, lokasi, PIC, catatan, dan lampiran.", "Simpan; gunakan Edit/Hapus sesuai izin."],
        expected="Agenda tampil pada daftar dan ringkasan terkait.",
        errors="Field wajib, status/tipe tidak valid, end sebelum start, atau file lampiran tidak sesuai.",
        risks="Penghapusan bersifat soft delete; pastikan agenda benar sebelum menghapus.",
        screenshot=None,
        notes="Source lengkap tersedia, tetapi halaman gagal dirender pada instance uji terisolasi karena query created_at ambigu. Perlu verifikasi melalui pengujian aplikasi pada lingkungan target."
    )
    feature_block(doc,
        title="Template Surat dan Builder Layout",
        purpose="Menstandarkan isi, nomor, ukuran kertas, margin, blok, signer, QR, dan mode dokumen.",
        access="Admin dan Sekretaris.",
        denied="Ketua, Bendahara, Bidang.",
        basis="templates.manage dan numbering.manage.",
        prerequisites="Kode/nama template, struktur layout, placeholder, dan konfigurasi nomor.",
        steps=["Buka Template Surat.", "Tambah atau edit template.", "Buka Builder untuk menyusun blok/layout.", "Simpan layout dan aktifkan template.", "Kelola profil penomoran bila diperlukan."],
        expected="Template dapat dipakai oleh surat baru.",
        errors="Kode duplikat, format/placeholder invalid, atau template masih direferensikan saat hapus.",
        risks="Perubahan template memengaruhi surat baru; uji preview sebelum dipakai luas.",
        screenshot="Gambar 13.",
        notes="Save layout template diverifikasi test otomatis."
    )
    feature_block(doc,
        title="Arsip Dokumen",
        purpose="Mengunggah, mencari, preview, dan mengunduh dokumen sekretariat.",
        access="Admin, Ketua, Sekretaris karena secretariat.view.",
        denied="Bendahara dan Bidang.",
        basis="secretariat.view untuk view dan store.",
        prerequisites="File valid dan metadata arsip.",
        steps=["Buka Arsip.", "Pilih Unggah Arsip.", "Isi metadata dan pilih file.", "Simpan.", "Gunakan Preview/Unduh."],
        expected="Record Document dan file tersimpan serta dapat ditelusuri.",
        errors="File hilang, tipe/ukuran ditolak, atau 404 saat storage record tidak cocok.",
        risks="Permission view juga mengizinkan upload pada route saat ini; batasi penggunaan sesuai SOP.",
        screenshot="Gambar 14.",
        notes="File lampiran tidak ikut ke backup database; hanya record database yang dibackup."
    )

    add_heading(doc, "6.7", "Anggota", 2)
    add_figure(doc, "07-members.png", "Gambar 15. Data Anggota.", ["Filter status/divisi/jabatan.", "Aksi tambah/edit/hapus sesuai permission.", "Detail member menggunakan data minimum operasional."])
    feature_block(doc,
        title="Data Anggota — Tambah, Ubah, Hapus",
        purpose="Mengelola identitas keanggotaan, kontak, demografi, pekerjaan, dan SIP.",
        access="Admin/Superadmin default.",
        denied="Ketua, Sekretaris, Bendahara, Bidang default.",
        basis="members.view/create/update/delete; pengaitan user login hanya Admin/Superadmin.",
        prerequisites="NPA, nama lengkap, status, dan master divisi/jabatan sesuai kebutuhan.",
        steps=["Buka Anggota > Data Anggota.", "Gunakan filter atau pencarian.", "Pilih Tambah Anggota atau Edit.", "Isi field dan simpan.", "Hapus hanya setelah memastikan tidak ada kebutuhan relasi aktif."],
        expected="Record member tersimpan/terbarui/soft deleted.",
        errors="NPA/email konflik, master data invalid, atau non-admin mencoba menautkan akun login.",
        risks="Perubahan divisi memengaruhi scope program kerja; penghapusan memengaruhi assignment dan iuran.",
        screenshot="Gambar 15.",
        notes="Test memastikan non-admin tidak dapat mengubah linked login account."
    )
    feature_block(doc,
        title="Import, Export, Template Import, dan Resolusi Konflik",
        purpose="Memproses data anggota massal dengan staging dan penyelesaian konflik.",
        access="Admin/Superadmin default.",
        denied="Role lain default.",
        basis="members.import, members.export, members.resolve_import.",
        prerequisites="Gunakan template unduhan; siapkan file dengan kolom dan nilai master yang benar.",
        steps=["Unduh Template Import.", "Isi data fiktif/terverifikasi sesuai format.", "Unggah dan jalankan Import.", "Buka daftar konflik batch.", "Pilih tindakan resolve per baris dan simpan.", "Gunakan Export untuk hasil terfilter."],
        expected="Baris valid masuk; konflik tercatat untuk diselesaikan; ekspor mengikuti filter.",
        errors="Format file/kolom invalid, duplicate NPA/email, target resolve tidak ditemukan, atau permission kurang.",
        risks="Import massal dapat menimpa/menduplikasi data bila resolusi salah; lakukan backup dan uji file kecil.",
        screenshot=None,
        notes="Perlu verifikasi manual dengan file contoh representatif dan data target."
    )

    add_heading(doc, "6.8", "Program Kerja", 2)
    add_figure(doc, "08-work-programs.png", "Gambar 16. Daftar Program Kerja.", ["Filter dan status.", "Aksi buat program sesuai permission.", "Program yang terlihat telah melalui scope backend."])
    add_figure(doc, "17-work-program-detail.png", "Gambar 17. Detail Program Kerja.", ["Metadata/status dan aksi workflow.", "Tab tugas/Gantt/monitoring/budget/dokumen/evaluasi.", "Aksi berubah sesuai status dan permission."])
    feature_block(doc,
        title="Daftar, Detail, Buat, Ubah, Hapus, dan Submit",
        purpose="Mencatat usulan program dan mengirimnya ke workflow review.",
        access="Admin/Ketua global; Sekretaris dalam scope untuk create/update/submit; Bendahara/Bidang lihat dalam scope.",
        denied="Pengguna di luar scope divisi, collaborator, primary PIC, assignment, task PIC/assignee.",
        basis="WorkProgramPolicy view/create/update/delete/submit; status draft atau revision_requested; visibleTo scope.",
        prerequisites="Periode, divisi, PIC, nama, tujuan/indikator, tanggal, dan anggaran valid.",
        steps=["Buka Daftar Program.", "Pilih Buat Program.", "Lengkapi informasi dasar, tujuan, indikator, jadwal, tim, dan anggaran.", "Simpan Draft.", "Edit bila perlu.", "Pilih Submit setelah data lengkap."],
        expected="Program draft tersimpan lalu berubah submitted; reviewer mendapat notifikasi sekali.",
        errors="Kode duplikat, tanggal terbalik, anggaran negatif, data submit belum lengkap, status terkunci, atau scope ditolak.",
        risks="Submit mengunci edit biasa sampai withdraw/revisi; hapus hanya untuk draft dan akses global.",
        screenshot="Gambar 16 dan 17.",
        notes="Ketua dapat create/update/submit dan juga review; self-review/self-approve tetap ditolak policy."
    )
    feature_block(doc,
        title="Workflow Review sampai Arsip",
        purpose="Mengendalikan perubahan status program secara berurutan dan terotorisasi.",
        access="Admin/Ketua untuk review/approve/reject/revision/archive; pelaksana scoped untuk submit, jadwal, progres, evaluasi sesuai permission.",
        denied="Submitter/creator tidak boleh approve/reject program sendiri; transisi ilegal ditolak.",
        basis="WorkProgram::TRANSITIONS dan WorkProgramPolicy.",
        prerequisites="Status sumber benar, alasan pada revisi/reject, serta permission dan scope.",
        steps=["Draft → Submit.", "Reviewer: Mulai Review.", "Pilih Minta Revisi, Approve, atau Reject.", "Program approved dijadwalkan.", "Mulai Pelaksanaan; Hold/Resume bila perlu.", "Selesaikan, isi Evaluasi, lalu Arsipkan."],
        expected="Audit approval mencatat setiap transisi dan aktor.",
        errors="Self-approval, alasan kosong, status sumber salah, duplicate approval, atau permission kurang.",
        risks="Approve/reject memengaruhi pelaksanaan; completed mengunci budget/dokumen sebelum evaluasi.",
        screenshot="Gambar 17.",
        notes="Completed dapat dibuka kembali ke revision_requested oleh reviewer."
    )
    feature_block(doc,
        title="Tugas, PIC, Assignee, Dependensi, Gantt, dan Collaborator Divisi",
        purpose="Menyusun rencana kerja rinci dan tanggung jawab lintas bidang.",
        access="Admin/Ketua/Sekretaris manage_tasks dalam scope; PIC/assignee dapat update progres tugas tertentu; Bendahara/Bidang lihat scoped.",
        denied="Pengguna tanpa view scope atau manage_tasks untuk perubahan struktur tugas.",
        basis="WorkProgramTaskPolicy; dependency types FS/SS/FF/SF; cycle validation.",
        prerequisites="Program belum terminal, tanggal tugas valid, PIC/assignee tersedia.",
        steps=["Buka tab Tugas/Gantt.", "Tambah tugas dan sub-tugas.", "Atur PIC, assignee, tanggal, bobot, milestone.", "Tambah dependensi dan collaborator divisi.", "Periksa Gantt dan jadwal massal.", "Update progres pada scheduled/in_progress."],
        expected="Dataset tugas/dependensi konsisten; notifikasi assignment tidak duplikat.",
        errors="Siklus hierarchy/dependency, milestone beda tanggal, task lintas program, atau lock_version konflik.",
        risks="Dependensi salah dapat memblokir jadwal; bulk schedule mengubah banyak tugas sekaligus.",
        screenshot="Gambar 17.",
        notes="PIC/assignee dapat update progres tanpa manage_tasks bila program berstatus scheduled/in_progress."
    )
    feature_block(doc,
        title="Progress, Budget, Dokumen, Risiko, Monitoring, Evaluasi, Audit Log, Laporan/Ekspor",
        purpose="Memantau realisasi, biaya, risiko, bukti, evaluasi, dan pelaporan program.",
        access="Sesuai permission update_progress/manage_budget/upload_document/evaluate/view_audit_log/export dan scope.",
        denied="Pengguna di luar scope; perubahan budget/dokumen pada status terminal.",
        basis="WorkProgramPolicy dan controller khusus monitoring/risks/budget/documents/evaluation/report.",
        prerequisites="Program terlihat; status mendukung aksi; file dan angka valid.",
        steps=["Buka tab Monitoring untuk metrik overdue/blocked/deadline.", "Update progres tugas.", "Kelola item budget dan realisasi.", "Unggah dokumen valid.", "Catat risiko/issue dan mitigasi.", "Pada completed, isi evaluasi.", "Gunakan Laporan dan ekspor PDF/CSV/XLSX/print bila diizinkan."],
        expected="Progres tertimbang, risiko, realisasi, dokumen, evaluasi, dan audit tersedia; ekspor mengikuti filter/scope.",
        errors="File MIME ditolak, budget terkunci, progres di luar 0–100, status salah, atau export permission kurang.",
        risks="Data finansial dan dokumen sensitif; pastikan scope serta klasifikasi sebelum upload/export.",
        screenshot="Gambar 17.",
        notes="Progres memakai leaf task; bila tanpa bobot memakai rata-rata. Test mencakup konsistensi ekspor dan scope."
    )

    add_heading(doc, "6.9", "Iuran", 2)
    add_figure(doc, "09-dues.png", "Gambar 18. Manajemen Iuran.", ["Pilih member dan periode.", "Input pembayaran sesuai role.", "Sinkronisasi/void mengikuti permission dan periode."])
    feature_block(doc,
        title="Pembayaran, Pembaruan, Sinkronisasi, dan Riwayat Member",
        purpose="Mencatat pembayaran iuran per periode dan menyinkronkan ledger anggota.",
        access="Admin; Bendahara kelola; Ketua lihat.",
        denied="Sekretaris dan Bidang default.",
        basis="dues.view/create/update/sync/manage; DuesPaymentPolicy; finance period open.",
        prerequisites="Member billable, periode belum dibayar, nominal/metode valid, periode keuangan terbuka.",
        steps=["Buka Iuran > Pembayaran.", "Cari member.", "Pilih periode awal dan durasi.", "Isi metode/referensi/catatan.", "Simpan pembayaran.", "Gunakan riwayat atau update sesuai izin."],
        expected="Payment, allocations per bulan, dan transaksi kas terkait dibuat atomik.",
        errors="Periode tumpang tindih, submit ganda sedang diproses, periode tutup, atau permission kurang.",
        risks="Pembayaran otomatis membuat transaksi kas yang tidak dapat diedit manual.",
        screenshot="Gambar 18.",
        notes="Durasi di atas 36 bulan diterima pada test; periksa kebutuhan bisnis sebelum input panjang."
    )
    feature_block(doc,
        title="Void, Rekap, dan Ekspor Iuran",
        purpose="Membatalkan pembayaran melalui kontrol approval dan membuat rekap.",
        access="Admin request/approve; Bendahara request void; Ketua approve void; rekap/ekspor default hanya Admin.",
        denied="Pemohon menyetujui sendiri; Bendahara tidak memiliki dues.recap.view/export default.",
        basis="dues.void.request/void/approve, dues.recap.view, dues.export, FinancialActionRequestPolicy.",
        prerequisites="Payment belum void, periode terbuka, alasan tersedia, approver berbeda.",
        steps=["Pilih pembayaran dan ajukan Void.", "Isi alasan.", "Approver membuka Audit & Approval.", "Approve atau Reject.", "Buka Rekap/Ekspor bila memiliki permission."],
        expected="Approved void menandai payment dan transaksi kas terkait; rekap mengecualikan void.",
        errors="Self-approval, request tidak pending, periode ditutup sebelum approval, atau permission kurang.",
        risks="Void mengubah laporan; simpan alasan dan bukti yang memadai.",
        screenshot="Gambar 18 dan Gambar 21.",
        notes="Menu Rekap Iuran tidak muncul untuk Bendahara dengan seeder default."
    )

    add_heading(doc, "6.10", "Kas / Transaksi", 2)
    add_figure(doc, "10-transactions.png", "Gambar 19. Transaksi Kas Masuk/Keluar.", ["Filter transaksi.", "Tambah transaksi sesuai permission.", "Edit metadata/void mengikuti status dan approval."])
    feature_block(doc,
        title="Tambah, Ubah, Lampiran, dan Void Transaksi",
        purpose="Mencatat pemasukan/pengeluaran serta jejak koreksi yang terkontrol.",
        access="Admin; Bendahara create/update/request void; Ketua lihat/approve void.",
        denied="Sekretaris dan Bidang default.",
        basis="transactions.view/create/update/void.request/void.approve/attachments.upload; CashTransactionPolicy.",
        prerequisites="Tanggal, tipe, kategori, metode, nominal positif, deskripsi; periode terbuka.",
        steps=["Buka Kas/Transaksi.", "Pilih Tambah Transaksi.", "Isi field dan lampiran bila tersedia.", "Simpan dan catat nomor transaksi otomatis.", "Edit hanya metadata non-finansial bila diizinkan.", "Ajukan Void; approver berbeda memutuskan di Audit."],
        expected="Transaksi posted bernomor unik; perubahan metadata diaudit; approved void dikeluarkan dari laporan.",
        errors="Field invalid, periode tutup, transaksi iuran, transaksi sudah void, submit duplikat, atau self-approval.",
        risks="Nilai finansial posted immutable; koreksi nominal harus melalui void dan transaksi pengganti, bukan edit langsung.",
        screenshot="Gambar 19.",
        notes="Permission transactions.update.metadata dan transactions.adjust.amount dibuat seeder tetapi route/policy saat ini tidak menggunakan keduanya."
    )

    add_heading(doc, "6.11", "Laporan dan Tutup Buku", 2)
    add_figure(doc, "11-report-cash.png", "Gambar 20. Laporan Kas.", ["Filter periode/kategori/metode.", "Kartu ringkasan dan grafik.", "Ekspor PDF memerlukan permission."])
    add_figure(doc, "12-financial-summary.png", "Gambar 21. Resume/Ikhtisar Keuangan.", ["Ringkasan kas dan iuran.", "Filter divisi/periode.", "PDF mengikuti permission export/print."])
    add_figure(doc, "13-finance-periods.png", "Gambar 22. Tutup Buku.", ["Daftar status periode.", "Role view hanya membaca.", "Tombol tutup hanya untuk finance.period.close."])
    feature_block(doc,
        title="Laporan Kas dan Resume Keuangan",
        purpose="Menganalisis saldo, arus kas, iuran tertagih/terkumpul, tren, top expense, dan tunggakan.",
        access="Admin/Superadmin default.",
        denied="Ketua, Sekretaris, Bendahara, Bidang tidak mendapat reports.* default.",
        basis="reports.cash.view, reports.financial.view, reports.export, reports.print.",
        prerequisites="Filter tanggal valid dan data ledger tersedia.",
        steps=["Buka Laporan Kas atau Resume Keuangan.", "Atur rentang tanggal/filter.", "Pilih Terapkan.", "Tinjau summary, grafik, dan detail.", "Pilih PDF bila memiliki export/print."],
        expected="Nilai summary konsisten dengan detail; void/soft delete tidak dihitung.",
        errors="403 tanpa permission; rentang/filter invalid; periode tanpa data menghasilkan empty state.",
        risks="Ekspor dapat memuat data finansial sensitif; simpan dan bagikan secara terkendali.",
        screenshot="Gambar 20 dan Gambar 21.",
        notes="Route legacy /reports, /reports/resume, /reports/export hanya auth+verified dan sebagian berisi halaman statis; gunakan route menu utama yang berpermission."
    )
    feature_block(doc,
        title="Tutup Buku / Periode Keuangan",
        purpose="Mengunci transaksi, iuran, void, dan perubahan metadata pada bulan yang telah ditutup.",
        access="Admin dapat lihat/tutup; Ketua dan Bendahara lihat.",
        denied="Sekretaris dan Bidang; Ketua/Bendahara tidak menutup default.",
        basis="finance.period.view/close/reopen; FinancePeriodPolicy; route close saja tersedia.",
        prerequisites="Tahun/bulan valid dan periode masih open.",
        steps=["Buka Laporan > Tutup Buku.", "Periksa bulan dan ringkasan.", "Admin pilih Tutup Periode.", "Konfirmasi setelah rekonsiliasi."],
        expected="Status closed; data historis tetap dapat dibaca tetapi mutasi ditolak.",
        errors="Permission kurang, bulan tidak valid, periode sudah closed.",
        risks="Setelah tutup, pending void juga tidak dapat disetujui; selesaikan rekonsiliasi dan approval dahulu.",
        screenshot="Gambar 22.",
        notes="Permission/policy reopen tersedia, tetapi tidak ditemukan route atau tombol reopen pada versi ini."
    )

    add_heading(doc, "6.12", "Audit & Approval", 2)
    add_figure(doc, "14-audit.png", "Gambar 23. Audit & Approval.", ["Daftar request pending/selesai.", "Filter aktivitas.", "Approve/Reject hanya untuk permission dan target sesuai."])
    feature_block(doc,
        title="Audit Aktivitas dan Approval Void",
        purpose="Menelusuri perubahan dan memutuskan permintaan void iuran/transaksi.",
        access="Admin dan Ketua lihat/approve; Bendahara lihat; role lain tidak default.",
        denied="Pemohon request yang sama, pengguna tanpa permission approve, request non-pending.",
        basis="activity.view, dues.void.approve, transactions.void.approve, FinancialActionRequestPolicy.",
        prerequisites="Request pending dan approver berbeda dari requester.",
        steps=["Buka Audit & Approval.", "Filter request pending.", "Buka detail target dan alasan.", "Pilih Approve atau Reject.", "Isi catatan bila diminta dan konfirmasi."],
        expected="Status request berubah; approve menjalankan void, reject tidak mengubah target.",
        errors="403 self-approval/permission; 422 periode tutup atau status request berubah.",
        risks="Approval mengubah laporan dan ledger; pastikan bukti memadai.",
        screenshot="Gambar 23.",
        notes="Backend diamankan walau tombol UI tersembunyi; diverifikasi test."
    )

    add_heading(doc, "6.13", "Pengaturan", 2)
    add_figure(doc, "15-settings.png", "Gambar 24. Pengaturan Admin.", ["Profil organisasi/branding.", "Master data dan iuran.", "User/role/permission, backup, dan reset dipisah dalam tab."])
    feature_block(doc,
        title="Profil Organisasi, Branding, Iuran, dan Master Data",
        purpose="Mengelola identitas aplikasi serta referensi operasional.",
        access="Role admin/superadmin pada group route; aksi form juga mengikuti validasi.",
        denied="Ketua, Sekretaris, Bendahara, Bidang meskipun settings.view individual diberikan.",
        basis="role:admin|superadmin; master data route berada dalam group yang sama.",
        prerequisites="Data yang disepakati organisasi dan logo non-sensitif.",
        steps=["Buka Pengaturan.", "Pilih tab Profil/Master Data/Pengaturan Iuran.", "Ubah nilai.", "Simpan dan verifikasi header, surat, serta laporan yang terkait."],
        expected="app_settings/master data/iuran diperbarui.",
        errors="File logo invalid, kode/nama duplikat, data dipakai relasi, atau role guard 403.",
        risks="Perubahan master data memengaruhi filter dan transaksi baru; jangan menghapus item yang masih digunakan.",
        screenshot="Gambar 24.",
        notes="app_settings pada database pengembangan yang diaudit kosong; fallback identitas tidak sepenuhnya konsisten antar halaman. Verifikasi nama/logo resmi sebelum publikasi."
    )
    feature_block(doc,
        title="Pengguna, Role, dan Permission",
        purpose="Membuat/menonaktifkan akun, reset password, assign role, serta sinkronisasi permission role/individual.",
        access="Admin/Superadmin yang juga memiliki permission users.*, roles.*, permissions.*.",
        denied="Role lain oleh route guard.",
        basis="FormRequest permission + role guard admin|superadmin; activity log.",
        prerequisites="Prinsip least privilege, pemilik akun teridentifikasi, dan rencana pemulihan akses.",
        steps=["Buka User & Permission.", "Buat atau pilih user.", "Tetapkan satu role utama.", "Tambahkan permission individual hanya bila benar-benar perlu.", "Uji menu, URL langsung, dan tombol.", "Catat pemilik keputusan dan tanggal review."],
        expected="Akses berubah dan audit log tercatat.",
        errors="Email/role/permission tidak ditemukan, permission operator kurang, atau proteksi superadmin.",
        risks="Perubahan salah dapat membuka data atau mengunci seluruh admin; pertahankan satu akun pemulihan superadmin yang aman.",
        screenshot="Gambar 24.",
        notes="Permission individual bersifat aditif; tidak ada deny eksplisit. Perhatikan bug perbandingan nama Admin/ admin yang dicatat pada Bab 10."
    )
    feature_block(doc,
        title="Backup dan Restore",
        purpose="Membuat dump data aplikasi dan mengembalikannya dari ZIP yang valid.",
        access="Admin/Superadmin dengan settings.view.",
        denied="Role lain.",
        basis="settings.backups.*; konfirmasi restore harus persis RESTORE DATABASE; ZIP maksimal 50 MB; test backup/restore lulus.",
        prerequisites="Waktu pemeliharaan, backup terbaru tervalidasi, dan pemilik bisnis menyetujui restore.",
        steps=["Pilih Backup Full Database.", "Unduh dan simpan file secara aman.", "Untuk restore, pilih ZIP hasil IDI Apps.", "Ketik frasa konfirmasi yang diminta.", "Jalankan restore dan lakukan rekonsiliasi pasca-restore."],
        expected="Record data kembali sesuai backup.",
        errors="ZIP bukan format IDI Apps, tabel inti tidak ada, ukuran lebih, atau proses restore gagal.",
        risks="Restore mengganti data. Backup hanya berisi record database; file lampiran fisik tidak disertakan.",
        screenshot="Gambar 24.",
        notes="Jangan melakukan restore tanpa salinan data saat ini dan verifikasi file lampiran terpisah."
    )
    feature_block(doc,
        title="Factory Reset — Hard, Finance, dan Custom",
        purpose="Mengosongkan data dalam scope tertentu.",
        access="Admin/Superadmin dengan settings.view.",
        denied="Semua role lain.",
        basis="settings.factory-reset.*; controller truncate tabel terpilih.",
        prerequisites="Persetujuan tertulis, backup tervalidasi, daftar tabel/scope, dan rencana pemulihan.",
        steps=["HENTIKAN operasi pengguna.", "Buat dan unduh backup.", "Verifikasi backup dan lampiran.", "Pilih scope reset yang paling sempit.", "Minta pemeriksa kedua meninjau scope.", "Jalankan reset hanya pada lingkungan yang benar.", "Segera ubah kredensial bootstrap dan rekonsiliasi data."],
        expected="Tabel dalam scope dikosongkan; hard reset membuat akun bootstrap dan me-seed ulang role/permission.",
        errors="Scope kosong atau tabel tidak terdaftar; ketergantungan data dapat meninggalkan record tersisa.",
        risks="SANGAT TINGGI dan sulit dipulihkan tanpa backup. Jangan jalankan sebagai prosedur harian.",
        screenshot="Gambar 24.",
        notes="Tidak dijalankan saat QA. Implementasi hard reset tidak meminta frasa konfirmasi, memakai nama tabel agenda yang tidak sesuai, dan belum mencantumkan tabel program kerja/organisasi terbaru. Akun bootstrap menggunakan kredensial yang dapat diprediksi pada source; nilainya sengaja tidak dicantumkan di buku ini."
    )

    add_heading(doc, "7", "Alur Kerja Lintas Role", 1)
    workflows = [
        ("Surat resmi", "Sekretaris membuat draft/layout → Sekretaris/Ketua finalisasi → signer yang ditetapkan menandatangani → publik memverifikasi → Sekretariat arsip/cabut sesuai permission."),
        ("Program kerja", "Sekretaris/Ketua membuat dan submit → Ketua/Admin review → revisi/approve/reject → PIC/assignee menjadwalkan dan melaksanakan → Bendahara/Ketua mengelola/meninjau budget sesuai akses → evaluator mengisi evaluasi → Ketua/Admin arsip."),
        ("Void iuran", "Bendahara/Admin mengajukan → Ketua/Admin lain meninjau → approve men-void payment dan transaksi terkait; reject mempertahankan data."),
        ("Void transaksi", "Bendahara/Admin mengajukan → Ketua/Admin lain meninjau → approved void dikeluarkan dari laporan."),
        ("Periode pengurus", "Sekretaris/Ketua/Admin membuat draft struktur dan assignment → publish → activate menyinkronkan akses → replace/end mencabut role terkelola dan menyimpan histori."),
        ("Tutup buku", "Bendahara rekonsiliasi → Ketua meninjau → Admin menutup periode setelah request approval selesai → semua role hanya membaca historis."),
    ]
    add_table(doc, ["Alur", "Urutan lintas role"], workflows, [1.35, 5.15], font_size=9)
    add_callout(doc, "Segregation of duties", "Requester void dan submitter/creator program tidak boleh menyetujui tindakan sendiri. Gunakan akun berbeda dan simpan alasan keputusan.", "ok")

    add_heading(doc, "8", "Status Workflow", 1)
    add_heading(doc, "8.1", "Program Kerja", 2)
    add_table(doc, ["Status", "Aksi masuk/keluar utama", "Batas edit"], [
        ("draft", "buat/edit → submit atau cancel", "Edit program; delete hanya draft + global view."),
        ("submitted", "withdraw atau start review", "Edit biasa ditolak."),
        ("under_review", "revision_requested / approved / rejected", "Reviewer berbeda dari creator/submitter."),
        ("revision_requested", "edit → resubmit atau cancel", "Creator/scoped editor dapat memperbaiki."),
        ("approved", "schedule atau cancel", "Tugas/budget masih dapat dikelola sesuai scope."),
        ("scheduled", "start execution atau cancel", "Progres tugas dapat diperbarui."),
        ("in_progress", "hold / complete / cancel", "Progres, task, budget, dokumen sesuai policy."),
        ("on_hold", "resume", "Tidak ada complete langsung."),
        ("completed", "revision_requested atau evaluated", "Budget/dokumen terkunci; evaluasi diizinkan."),
        ("evaluated", "archive", "Operasional terkunci."),
        ("archived/rejected/cancelled", "terminal", "Tidak dapat dikelola normal."),
    ], [1.25, 2.6, 2.65], font_size=8.7)
    add_heading(doc, "8.2", "Status Lain", 2)
    add_table(doc, ["Objek", "Status", "Aturan penting"], [
        ("Tugas", "todo, in_progress, blocked, completed, cancelled", "Completed wajib progres 100; milestone satu tanggal."),
        ("Periode organisasi", "draft, published, active, ended, archived", "Hanya satu active; ended/archived read-only."),
        ("Assignment", "draft, active, replaced, ended, cancelled", "Draft belum memberi akses; replace/end menjaga histori."),
        ("Periode keuangan", "open, closed", "Closed menolak transaksi, iuran, void, dan edit metadata."),
        ("Action request", "pending, approved, rejected", "Hanya pending dan bukan requester dapat direview."),
        ("Agenda", "planned, done, cancelled", "Tipe internal/external."),
        ("Surat", "draft, finalized, archived + is_revoked", "PDF/arsip butuh final; cabut mengubah validitas."),
    ], [1.35, 2.15, 3.0], font_size=8.8)

    add_heading(doc, "9", "Tindakan Berisiko dan Kontrol Persetujuan", 1)
    risk_rows = [
        ("Hard/custom/finance reset", "Kehilangan data, scope tidak lengkap", "Backup; pemeriksa kedua; hentikan operasi; rekonsiliasi; jangan gunakan tanpa perbaikan catatan implementasi."),
        ("Restore", "Menimpa data dan record file", "Validasi sumber ZIP; backup kondisi saat ini; simpan lampiran terpisah; uji hasil."),
        ("Role/permission", "Privilege escalation atau lockout", "Least privilege; pertahankan superadmin pemulihan; uji menu+URL; review berkala."),
        ("Void", "Mengubah ledger/laporan", "Requester ≠ approver; alasan dan bukti; selesaikan sebelum tutup buku."),
        ("Finalisasi/cabut surat", "Nomor resmi/keabsahan berubah", "Preview; reviewer; cek signer; alasan pencabutan; verifikasi publik."),
        ("Tutup buku", "Mengunci koreksi", "Rekonsiliasi, selesaikan pending approval, persetujuan Ketua, baru Admin close."),
        ("Import anggota", "Duplikasi/overwrite massal", "Template resmi; backup; file kecil; review konflik; sampling hasil."),
        ("Aktivasi periode pengurus", "Role/divisi akun berubah", "Preflight; review posisi wajib; cek akun/member/role; audit setelah aktivasi."),
    ]
    add_table(doc, ["Tindakan", "Risiko", "Kontrol minimum"], risk_rows, [1.45, 1.75, 3.3], font_size=8.7)

    add_heading(doc, "10", "Catatan Implementasi yang Perlu Ditindaklanjuti", 1)
    impl_notes = [
        ("I-01", "Dashboard", "Semua akun verified dapat melihat metrik finansial; quick action tetap tampak untuk role yang akhirnya menerima 403.", "Tinggi"),
        ("I-02", "Pengaturan", "Ketua/Bendahara memiliki settings.view dan melihat menu, tetapi route mensyaratkan role admin|superadmin sehingga 403.", "Sedang"),
        ("I-03", "Tanda Tangan", "Route inbox dan sign tidak memiliki middleware permission khusus; controller sign membatasi kepemilikan signer dan status.", "Tinggi"),
        ("I-04", "Role Admin", "Pemeriksaan literal “Admin” tidak cocok dengan role seeded “admin” untuk proteksi hapus/self-demotion.", "Tinggi"),
        ("I-05", "Reset", "Hard reset tanpa frasa konfirmasi; daftar tabel memakai agenda (bukan agendas) dan belum mencakup tabel program kerja/organisasi terbaru.", "Kritis"),
        ("I-06", "Backup", "ZIP backup berisi dump record database, bukan file lampiran fisik.", "Tinggi"),
        ("I-07", "Tutup Buku", "finance.period.reopen ada di seeder/policy tetapi tidak ada route/tombol.", "Sedang"),
        ("I-08", "Permission transaksi", "transactions.update.metadata dan transactions.adjust.amount dibuat tetapi tidak digunakan oleh route/policy aktif.", "Rendah"),
        ("I-09", "Legacy reports", "Beberapa route /reports legacy hanya auth+verified dan controller halaman bersifat statis/placeholder.", "Sedang"),
        ("I-10", "Agenda", "Instance uji terisolasi gagal merender eager-load lampiran karena created_at ambigu; perlu uji ulang pada lingkungan target.", "Sedang"),
        ("I-11", "Identitas", "app_settings kosong pada database yang diaudit dan fallback nama organisasi berbeda antar halaman.", "Sedang"),
    ]
    add_table(doc, ["ID", "Area", "Temuan", "Prioritas"], impl_notes, [0.5, 1.1, 4.15, 0.75], font_size=8.2)

    add_heading(doc, "11", "Troubleshooting dan FAQ", 1)
    faqs = [
        ("Menu tidak terlihat", "Pastikan role/permission default atau individual tersedia. Refresh sesi setelah perubahan. Tetap uji URL langsung."),
        ("Menu terlihat tetapi 403", "UI dan route dapat memakai syarat berbeda. Contoh saat ini: Pengaturan pada Ketua/Bendahara. Hubungi Admin; jangan menambah permission tanpa analisis role guard."),
        ("Login selalu gagal", "Periksa email, kata sandi, status aktif, dan rate limit. Admin dapat reset password bila berizin."),
        ("Diminta verifikasi email", "Selesaikan tautan verifikasi; seluruh modul internal utama memakai middleware verified."),
        ("Tidak dapat edit surat", "Periksa status, letters.create/update, dan apakah surat sudah final/archived."),
        ("Nomor surat berubah saat finalisasi", "Generate hanya pratinjau; finalisasi mengunci nomor pertama yang belum dipakai."),
        ("Tidak dapat tanda tangan", "Akun harus terhubung ke member signer, surat final, dan tidak dicabut."),
        ("Import anggota konflik", "Buka batch konflik dan gunakan resolve_import; jangan mengulang unggah tanpa meninjau hasil."),
        ("Program kerja tidak terlihat", "Periksa divisi member, collaborator division, primary PIC, assignment, task PIC/assignee, dan permission view_own_field."),
        ("Program tidak dapat diedit", "Hanya draft/revision_requested dan pengguna dalam scope dengan work_program.update."),
        ("Approve program ditolak", "Status harus under_review; creator/submitter tidak boleh self-approve."),
        ("Progress tugas ditolak", "Program harus scheduled/in_progress dan pengguna harus punya update_progress atau menjadi PIC/assignee."),
        ("Pembayaran iuran duplikat", "Periode member sudah dialokasikan atau submit masih diproses. Jangan submit ulang."),
        ("Transaksi tidak bisa diubah", "Field finansial posted immutable; transaksi dari iuran tidak dapat diedit manual."),
        ("Void ditolak", "Periksa request masih pending, approver bukan requester, target belum void, dan periode masih open."),
        ("Laporan tidak memuat transaksi", "Void/soft delete dikeluarkan; cek tanggal, kategori, metode, dan timezone aplikasi."),
        ("Tidak bisa tutup buku", "Hanya finance.period.close; role Ketua/Bendahara default hanya view."),
        ("Tidak ada tombol buka kembali periode", "Belum ada route/tombol reopen pada versi saat ini; jangan mengubah data di luar aplikasi."),
        ("Restore gagal", "Pastikan ZIP berasal dari backup IDI Apps, ukurannya sesuai, dan memuat tabel inti. Eskalasi sebelum mencoba ulang."),
        ("Lampiran hilang setelah restore", "Backup data tidak menyertakan file fisik. Pulihkan file dari salinan terpisah dan rekonsiliasi record."),
    ]
    for idx, (q, a) in enumerate(faqs, 1):
        add_para(doc, f"{idx}. {q}", bold=True, color=NAVY, size=10.5, before=5, after=2, keep=True)
        add_para(doc, a, size=10.2, after=5)

    add_heading(doc, "12", "Fitur yang Perlu Verifikasi Manual", 1)
    manual = [
        "Pengiriman email reset password, verifikasi email, dan reset link akun yang dibuat melalui assignment organisasi.",
        "Halaman Agenda pada lingkungan target serta upload/preview lampiran agenda.",
        "Kualitas PDF surat pada template organisasi resmi, header image, QR multi-signer, arsip, dan skenario surat dicabut.",
        "Import/export anggota menggunakan file produksi yang sudah dianonimkan, termasuk seluruh pilihan resolusi konflik.",
        "Restore dan semua varian reset setelah temuan implementasi diperbaiki serta diuji di lingkungan non-produksi.",
        "Kesesuaian nama/logo/owner resmi karena app_settings saat audit belum berisi profil organisasi.",
        "Seluruh transisi organisasi (publish/activate/end/replace) dengan struktur dan pemilik keputusan nyata.",
        "Aksi reopen periode keuangan setelah route/UI tersedia.",
        "Uji keamanan tambahan untuk route Dashboard, Tanda Tangan, legacy reports, dan kontrol self-demotion/hapus role Admin.",
    ]
    add_bullets(doc, manual)

    add_heading(doc, "13", "Lampiran A — Daftar Permission Default", 1)
    add_para(doc, f"Seeder mendefinisikan {len(permissions)} permission. Tabel berikut menunjukkan role default yang mewarisinya; Bidang dipetakan ke role anggota.")
    groups = {}
    for perm in permissions:
        prefix = perm.split(".")[0]
        groups.setdefault(prefix, []).append(perm)
    for prefix, perms in groups.items():
        add_para(doc, prefix.replace("_", " ").title(), bold=True, color=DARK_RED, size=11, before=8, after=4, keep=True)
        rows = []
        for perm in perms:
            rows.append((
                perm,
                "✓" if perm in role_map["admin"] else "—",
                "✓" if perm in role_map["ketua"] else "—",
                "✓" if perm in role_map["sekretaris"] else "—",
                "✓" if perm in role_map["bendahara"] else "—",
                "✓" if perm in role_map["anggota"] else "—",
            ))
        add_table(doc, ["Permission", "Admin", "Ketua", "Sekr.", "Bend.", "Bidang"], rows,
                  [3.35, 0.63, 0.63, 0.63, 0.63, 0.63], font_size=7.8)

    add_heading(doc, "14", "Lampiran B — Kamus Istilah", 1)
    glossary = [
        ("Approval", "Keputusan approve/reject oleh pengguna yang berwenang dan berbeda dari pemohon bila policy mensyaratkan."),
        ("Assignee", "Pengguna yang ditugaskan pada task program kerja."),
        ("Bidang", "Unit/divisi kerja; dalam portal dipetakan aman ke anggota + scope, bukan role literal."),
        ("Builder", "Penyusun layout/blok surat atau template."),
        ("Collaborator division", "Divisi tambahan yang memperoleh scope lihat/kolaborasi program."),
        ("Direct permission", "Permission individual yang ditambahkan langsung pada user."),
        ("Finalisasi", "Penguncian nomor dan pembuatan versi/PDF surat final."),
        ("Gantt", "Visual jadwal tugas dan dependensi program."),
        ("Ledger", "Catatan transaksi/alokasi yang menjadi dasar laporan."),
        ("PIC", "Penanggung jawab utama program atau tugas."),
        ("Policy", "Aturan backend yang menggabungkan permission, status, kepemilikan, dan scope."),
        ("Public hash", "Token publik untuk verifikasi dokumen tanpa login."),
        ("Role permission", "Permission yang diwariskan dari role."),
        ("Scope", "Batas data berdasarkan divisi, collaborator, PIC, assignment, atau assignee."),
        ("Soft delete", "Record ditandai terhapus tanpa langsung dihapus fisik."),
        ("Void", "Pembatalan finansial yang mempertahankan jejak record."),
        ("Workflow", "Urutan status dan aksi yang diizinkan."),
    ]
    add_table(doc, ["Istilah", "Definisi"], glossary, [1.45, 5.05], font_size=9)

    add_heading(doc, "15", "Lampiran C — Catatan Role Bidang", 1)
    add_callout(doc, "Kesimpulan", "Tidak ada role literal bidang. Implementasi yang paling aman untuk istilah pengguna adalah “Bidang (profil kerja)” yang terdiri dari role anggota, keterkaitan Member.division_id, dan scope tambahan dari primary PIC, assignment program, task PIC/assignee, atau collaborator division.", "note")
    add_bullets(doc, [
        "Master divisi memuat nama-nama Bidang organisasi.",
        "Master posisi memuat jabatan Ketua Bidang, Sekretaris Bidang, Koordinator, dan Anggota Pengurus.",
        "Assignment pengurus memilih unit, slot jabatan, member, dan portal_role_id secara terpisah.",
        "Role anggota default: work_program.view_own_field, work_program.update_progress, work_program.upload_document.",
        "Menyebutnya sebagai role Bidang akan menyesatkan karena Admin dapat memilih role portal lain untuk assignment organisasi.",
    ])

    add_heading(doc, "16", "Lampiran D — Checklist QA Dokumen", 1)
    add_bullets(doc, [
        "Semua menu utama pada Menu.jsx telah dipetakan.",
        "Seluruh 92 permission default seeder dimuat pada Lampiran A.",
        "Role literal dan profil Bidang dijelaskan tanpa mengarang role baru.",
        "Aturan status program, organisasi, finansial, agenda, surat, dan approval dicatat.",
        "Perbedaan UI/route/backend diberi label Catatan Implementasi.",
        "Screenshot menggunakan data fiktif dan tidak menampilkan kredensial.",
        "217 test / 1.263 assertion lulus pada tanggal audit.",
        "Item yang belum dapat diuji penuh dicantumkan pada Bab 12.",
    ])
    # Reapply headers/footers and section geometry after all sections exist.
    for i, section in enumerate(doc.sections):
        if i == 1:
            configure_section(section, landscape=True)
            add_running_header_footer(section, "IDI Apps • Matriks Akses")
        else:
            configure_section(section, landscape=False)
            add_running_header_footer(section)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
